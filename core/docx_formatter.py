"""
docx_formatter.py
Минимальный модуль для замены текста в DOCX с сохранением форматирования
"""

class DocxFormatter:
    """Замена текста в ячейке с сохранением форматирования первого run"""

    @staticmethod
    def replace_text(cell, new_text, preserve_formatting=True):
        """Замена текста в ячейке с защитой от дублирования"""
        if not preserve_formatting:
            # Полная очистка и простая замена
            for p in cell.paragraphs:
                p.clear() if hasattr(p, "clear") else None
            if not cell.paragraphs:
                p = cell.add_paragraph()
                p.add_run(new_text)
            else:
                cell.paragraphs[0].add_run(new_text)
            return

        # Получаем текущие параграфы
        paragraphs = cell.paragraphs

        # Если нет параграфов, создаем новый
        if not paragraphs:
            p = cell.add_paragraph()
            p.add_run(new_text)
            return

        # Берем первый параграф для форматирования
        p = paragraphs[0]

        # Если параграф пустой, добавляем run
        if not p.runs:
            p.add_run(new_text)
            return

        # Сохраняем форматирование первого run
        first_run = p.runs[0]
        formatting = {
            "bold": first_run.bold,
            "italic": first_run.italic,
            "underline": first_run.underline,
            "font_name": first_run.font.name,
            "font_size": first_run.font.size,
            "color": first_run.font.color.rgb if first_run.font.color else None
        }

        # ВАЖНО: Полностью очищаем ячейку от лишних параграфов
        while len(cell.paragraphs) > 1:
            last_p = cell.paragraphs[-1]
            last_p._p.getparent().remove(last_p._p)

        # Очищаем оставшийся параграф
        p.clear()

        # Создаем новый run с текстом
        new_run = p.add_run(new_text)

        # Применяем форматирование
        new_run.bold = formatting["bold"]
        new_run.italic = formatting["italic"]
        new_run.underline = formatting["underline"]

        if formatting["font_name"]:
            new_run.font.name = formatting["font_name"]

        if formatting["font_size"]:
            new_run.font.size = formatting["font_size"]

        if formatting["color"]:
            new_run.font.color.rgb = formatting["color"]

    @staticmethod
    def replace_with_preserve_special_formatting(cell, new_text, original_run):
        """
        Заменяет текст с сохранением специального форматирования
        (цвет, подчеркивание, и т.д.)
        """
        # Сохраняем все атрибуты форматирования
        formatting = {
            "bold": original_run.bold,
            "italic": original_run.italic,
            "underline": original_run.underline,
            "font_name": original_run.font.name,
            "font_size": original_run.font.size,
            "color": original_run.font.color.rgb if original_run.font.color else None,
        }

        # Очищаем ячейку
        for p in cell.paragraphs:
            p.clear()

        if not cell.paragraphs:
            p = cell.add_paragraph()
        else:
            p = cell.paragraphs[0]

        # Создаем новый run с текстом
        new_run = p.add_run(new_text)

        # Применяем все сохраненное форматирование
        new_run.bold = formatting["bold"]
        new_run.italic = formatting["italic"]
        new_run.underline = formatting["underline"]

        if formatting["font_name"]:
            new_run.font.name = formatting["font_name"]

        if formatting["font_size"]:
            new_run.font.size = formatting["font_size"]

        if formatting["color"]:
            new_run.font.color.rgb = formatting["color"]

    @staticmethod
    def replace_text_with_background(cell, new_text, background_color=None):
        """
        Заменяет текст в ячейке с сохранением цветного фона

        Args:
            background_color: цвет фона в hex (например, "FFFF00" для желтого)
        """
        if not cell.paragraphs:
            p = cell.add_paragraph()
            run = p.add_run(new_text)
        else:
            p = cell.paragraphs[0]
            if not p.runs:
                run = p.add_run(new_text)
            else:
                # Используем форматирование первого run
                first_run = p.runs[0]

                # Очищаем параграф
                p._p.clear_content()

                # Создаем новый run
                run = p.add_run(new_text)

                # Копируем форматирование
                run.bold = first_run.bold
                run.italic = first_run.italic
                run.underline = first_run.underline

                if first_run.font.name:
                    run.font.name = first_run.font.name

                if first_run.font.size:
                    run.font.size = first_run.font.size

                if first_run.font.color:
                    run.font.color.rgb = first_run.font.color.rgb

        # Применяем цвет фона если указан
        if background_color:
            from docx.oxml import OxmlElement
            rPr = run._r.get_or_add_rPr()
            shd = OxmlElement('w:shd')
            shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', background_color)
            rPr.append(shd)