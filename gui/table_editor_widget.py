import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from docx import Document

from core.table_template import TableTemplate
from core.docx_formatter import DocxFormatter


class TableEditorWidget(ttk.Frame):
    """
    Встраиваемый редактор таблиц:
    - не создаёт своё окно
    - работает с переданным Document
    - адаптивно подстраивается под размер
    """

    def __init__(self, parent):
        super().__init__(parent)

        self.document = None
        self.current_table = None
        self.current_file_path = None
        self.selected_cells = set()
        self.preserve_formatting = True
        self.template = None

        self.is_modified = False  # индикатор изменений

        self._build_ui()
        self.bind("<Configure>", self._on_resize)

    # ---------------- UI ----------------

    def _build_ui(self):
        self.main_pane = ttk.PanedWindow(self, orient=tk.HORIZONTAL)
        self.main_pane.pack(fill=tk.BOTH, expand=True)

        self.left = ttk.Frame(self.main_pane)
        self.main_pane.add(self.left, weight=3)

        self.right = ttk.Frame(self.main_pane)
        self.main_pane.add(self.right, weight=2)

        self.table_combo = ttk.Combobox(self.left, state="readonly")
        self.table_combo.pack(fill=tk.X, pady=2)
        self.table_combo.bind("<<ComboboxSelected>>", self.on_table_select)

        self.canvas = tk.Canvas(self.left)
        self.canvas.pack(fill=tk.BOTH, expand=True)

        self.table_frame = ttk.Frame(self.canvas)
        self.canvas_window = self.canvas.create_window((0, 0), window=self.table_frame, anchor="nw")

        self.table_frame.bind(
            "<Configure>",
            lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all"))
        )

        # ----------- правая часть -----------

        edit = ttk.LabelFrame(self.right, text="Редактирование ячейки")
        edit.pack(fill=tk.X, pady=5)

        ttk.Label(edit, text="Текущий текст:").pack(anchor="w")
        self.old_text = tk.Text(edit, height=2, state=tk.DISABLED, font=("Arial", 9))
        self.old_text.pack(fill=tk.X)

        ttk.Label(edit, text="Новый текст:").pack(anchor="w", pady=(4, 0))
        self.new_text = tk.Text(edit, height=2, font=("Arial", 9))
        self.new_text.pack(fill=tk.X)

        ttk.Button(edit, text="Заменить", command=self.replace_cell).pack(fill=tk.X, pady=4)

        self.format_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(
            edit,
            text="Сохранять форматирование",
            variable=self.format_var
        ).pack(anchor="w")

        # Кнопка для работы с фоном
        ttk.Button(edit, text="Замена с сохранением фона", command=self.replace_with_background).pack(fill=tk.X, pady=4)

        # Индикатор изменений
        self.modified_label = ttk.Label(edit, text="Документ изменён", foreground="red")
        self.modified_label.pack(anchor="w")
        self.modified_label.pack_forget()

        tmpl = ttk.LabelFrame(self.right, text="Шаблоны таблиц")
        tmpl.pack(fill=tk.BOTH, expand=True, pady=5)

        ttk.Button(tmpl, text="Создать шаблон из текущей таблицы", command=self.create_template).pack(fill=tk.X, pady=2)
        ttk.Button(tmpl, text="Загрузить шаблон", command=self.load_template).pack(fill=tk.X, pady=2)
        ttk.Button(tmpl, text="Заполнить таблицу по шаблону", command=self.fill_table).pack(fill=tk.X, pady=2)

        # Кнопка сохранения документа
        ttk.Button(tmpl, text="Сохранить документ", command=self.save_document).pack(fill=tk.X, pady=2)

    # ---------------- Адаптивность ----------------

    def _on_resize(self, event):
        # Панель не поддерживает изменение ориентации после создания.
        pass

    # ---------------- Публичные методы ----------------

    def load_document(self, document: Document, file_path=None):
        self.document = document
        self.current_file_path = file_path
        self.is_modified = False
        self.modified_label.pack_forget()

        if not self.document.tables:
            self.table_combo["values"] = []
            self.current_table = None
            self._redraw_table()
            return

        self.table_combo["values"] = [f"Таблица {i + 1}" for i in range(len(self.document.tables))]
        self.table_combo.current(0)
        self.on_table_select()

    def save_document(self):
        if not self.document or not self.current_file_path:
            messagebox.showwarning("Нет документа", "Документ не загружен.")
            return

        try:
            self.document.save(self.current_file_path)
            self.is_modified = False
            self.modified_label.pack_forget()
            messagebox.showinfo("Сохранено", f"Документ сохранён:\n{self.current_file_path}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить документ:\n{e}")

    # ---------------- Логика ----------------

    def on_table_select(self, event=None):
        if not self.document:
            return

        idx = self.table_combo.current()
        if idx < 0:
            return

        self.current_table = self.document.tables[idx]
        self.selected_cells.clear()
        self._redraw_table()

    def _redraw_table(self):
        for w in self.table_frame.winfo_children():
            w.destroy()

        if not self.current_table:
            return

        for r, row in enumerate(self.current_table.rows):
            for c, cell in enumerate(row.cells):
                # Проверяем, есть ли фон в ячейке
                has_background = False
                for p in cell.paragraphs:
                    for run in p.runs:
                        rPr = run._r.rPr
                        if rPr is not None:
                            shd = rPr.find('.//w:shd', namespaces={
                                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                            if shd is not None:
                                fill = shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                                if fill and fill.upper() != 'FFFFFF' and fill.upper() != 'AUTO':
                                    has_background = True
                                    break
                    if has_background:
                        break

                # Устанавливаем цвет фона в зависимости от наличия фона в тексте
                bg_color = "lightyellow" if has_background else "white"
                if (r, c) in self.selected_cells:
                    bg_color = "lightblue"

                frame = tk.Frame(
                    self.table_frame,
                    bd=2 if (r, c) in self.selected_cells else 1,
                    relief=tk.SOLID,
                    bg=bg_color
                )
                frame.grid(row=r, column=c, sticky="nsew")

                text_widget = tk.Text(
                    frame,
                    height=2,
                    width=14,
                    wrap="word",
                    font=("Arial", 9),
                    state=tk.NORMAL
                )
                text_widget.insert("1.0", cell.text)
                text_widget.config(state=tk.DISABLED)
                text_widget.pack(fill=tk.BOTH, expand=True)

                frame.bind("<Button-1>", lambda e, rr=r, cc=c: self.on_cell_click(rr, cc))
                text_widget.bind("<Button-1>", lambda e, rr=r, cc=c: self.on_cell_click(rr, cc))

        for i in range(len(self.current_table.rows)):
            self.table_frame.rowconfigure(i, weight=1)
        if self.current_table.rows:
            for i in range(len(self.current_table.rows[0].cells)):
                self.table_frame.columnconfigure(i, weight=1)

        self.canvas.configure(scrollregion=self.canvas.bbox("all"))

    def on_cell_click(self, r, c):
        self.selected_cells = {(r, c)}

        cell = self.current_table.rows[r].cells[c]

        self.old_text.config(state=tk.NORMAL)
        self.old_text.delete("1.0", tk.END)
        self.old_text.insert("1.0", cell.text)
        self.old_text.config(state=tk.DISABLED)

        self.new_text.delete("1.0", tk.END)

        self._redraw_table()

    def replace_cell(self):
        if not self.selected_cells or not self.current_table:
            return

        r, c = next(iter(self.selected_cells))
        new_value = self.new_text.get("1.0", tk.END).strip()

        if not new_value:
            return

        DocxFormatter.replace_text(
            self.current_table.rows[r].cells[c],
            new_value,
            preserve_formatting=self.format_var.get()
        )

        self.is_modified = True
        self.modified_label.pack(anchor="w")

        self._redraw_table()

    def replace_with_background(self):
        """Замена текста с сохранением фона"""
        if not self.selected_cells or not self.current_table:
            return

        r, c = next(iter(self.selected_cells))
        new_value = self.new_text.get("1.0", tk.END).strip()

        if not new_value:
            return

        # Определяем, есть ли фон в исходной ячейке
        cell = self.current_table.rows[r].cells[c]
        has_background = False
        background_color = None

        for p in cell.paragraphs:
            for run in p.runs:
                rPr = run._r.rPr
                if rPr is not None:
                    shd = rPr.find('.//w:shd',
                                   namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                    if shd is not None:
                        fill = shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                        if fill and fill.upper() != 'FFFFFF' and fill.upper() != 'AUTO':
                            has_background = True
                            background_color = fill
                            break
            if has_background:
                break

        if has_background:
            # Используем метод с сохранением фона
            DocxFormatter.replace_text_with_background(cell, new_value, background_color)
        else:
            # Обычная замена
            DocxFormatter.replace_text(cell, new_value, preserve_formatting=self.format_var.get())

        self.is_modified = True
        self.modified_label.pack(anchor="w")

        self._redraw_table()

    # -------- шаблоны --------

    def load_template_from_data(self, template):
        self.template = template

    def create_template(self):
        if not self.current_table:
            messagebox.showwarning("Нет таблицы", "Сначала выберите таблицу.")
            return

        tmpl = TableTemplate.create_template(self.current_table)
        path = filedialog.asksaveasfilename(defaultextension=".json", filetypes=[("JSON", "*.json")])
        if not path:
            return

        TableTemplate.save_template(tmpl, path)
        messagebox.showinfo("OK", "Шаблон сохранён")

    def load_template(self):
        path = filedialog.askopenfilename(filetypes=[("JSON", "*.json")])
        if not path:
            return

        self.template = TableTemplate.load_template(path)
        messagebox.showinfo("OK", "Шаблон загружен")

    def fill_table(self):
        if not self.template:
            messagebox.showwarning("Нет шаблона", "Сначала загрузите шаблон.")
            return

        if not self.current_table:
            messagebox.showwarning("Нет таблицы", "Сначала выберите таблицу.")
            return

        TableTemplate.fill_table(
            self.current_table,
            self.template,
            data_map={},
            formatter=DocxFormatter
        )

        self.is_modified = True
        self.modified_label.pack(anchor="w")

        self._redraw_table()
        messagebox.showinfo("Готово", "Таблица заполнена из шаблона.")