import os
import shutil
import tempfile
import logging
import json
import datetime
import re
import functools
from contextlib import contextmanager
from typing import Dict, Any, Optional, List, Tuple
from collections import OrderedDict

from docx import Document
from docx.shared import Inches, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from PIL import Image

from .config import SignatureConfig
from .pdf_converter import PdfConverter
from .docx_formatter import DocxFormatter


def log_execution(func):
    """Декоратор для логирования выполнения методов"""

    @functools.wraps(func)
    def wrapper(*args, **kwargs):
        logging.info(f"Начало: {func.__name__}")
        try:
            result = func(*args, **kwargs)
            logging.info(f"Завершено: {func.__name__}")
            return result
        except Exception as e:
            logging.error(f"Ошибка в {func.__name__}: {e}")
            raise

    return wrapper


class DocxProcessor:
    SIGNATURE_PLACEHOLDER = "{sign}"
    STAMP_PLACEHOLDER = "{stamp}"
    DIRECTOR_PLACEHOLDER = "{director}"
    FULLNAME_PLACEHOLDER = "{FULL_NAME_AND_POSITION}"
    COMPANY_LOGO_PLACEHOLDER = "{company_logo}"
    BRACKET_PLACEHOLDER_PATTERN = r'\[([\s\S]+?)\]'

    # Константы
    EMUS_PER_INCH = 914400  # EMU в дюйме
    MAX_LOGO_WIDTH = 7.5  # дюймов
    LOGO_MARGIN = 0.5  # дюймов
    DEFAULT_IMAGE_WIDTH = 2.0  # дюймов
    MAX_CACHE_SIZE = 100  # максимальный размер кэша изображений

    def __init__(self, signature_config=None, data_dir=None):
        self.signature_config = signature_config or SignatureConfig()
        self.converter = PdfConverter()
        self.data_dir = data_dir
        self._image_cache = OrderedDict()
        self._regex_replacements = []
        self._history = []  # История изменений

        # Настройки для изображений подписи и печати
        self.signature_image_params = {
            "path": None,
            "width": 30,  # мм
            "height": 30,  # мм
            "offset_x": 0,  # мм
            "offset_y": -20  # мм
        }

        self.stamp_image_params = {
            "path": None,
            "width": 30,  # мм
            "height": 30,  # мм
            "offset_x": 0,  # мм
            "offset_y": -20  # мм
        }

    # ---------------------------------------------------------
    # Настройка параметров изображений
    # ---------------------------------------------------------
    def set_signature_image_params(self, path=None, width=None, height=None,
                                   offset_x=None, offset_y=None):
        """Устанавливает параметры для изображения подписи"""
        if path is not None:
            self.signature_image_params["path"] = path
        if width is not None:
            self.signature_image_params["width"] = width
        if height is not None:
            self.signature_image_params["height"] = height
        if offset_x is not None:
            self.signature_image_params["offset_x"] = offset_x
        if offset_y is not None:
            self.signature_image_params["offset_y"] = offset_y

    def set_stamp_image_params(self, path=None, width=None, height=None,
                               offset_x=None, offset_y=None):
        """Устанавливает параметры для изображения печати"""
        if path is not None:
            self.stamp_image_params["path"] = path
        if width is not None:
            self.stamp_image_params["width"] = width
        if height is not None:
            self.stamp_image_params["height"] = height
        if offset_x is not None:
            self.stamp_image_params["offset_x"] = offset_x
        if offset_y is not None:
            self.stamp_image_params["offset_y"] = offset_y

    def set_company(self, company_name):
        """Устанавливает компанию и автоматически загружает пути к изображениям"""
        if not self.data_dir:
            logging.warning("data_dir не указан, автоматическая загрузка изображений невозможна")
            return

        company_dir = os.path.join(self.data_dir, "organizations", company_name)
        if not os.path.isdir(company_dir):
            logging.warning(f"Папка компании не найдена: {company_dir}")
            return

        # Автоматически загружаем пути к изображениям
        signature_path = os.path.join(company_dir, "sign.png")
        if os.path.exists(signature_path):
            self.signature_image_params["path"] = signature_path
            logging.info(f"Загружена подпись: {signature_path}")

        stamp_path = os.path.join(company_dir, "stamp.png")
        if os.path.exists(stamp_path):
            self.stamp_image_params["path"] = stamp_path
            logging.info(f"Загружена печать: {stamp_path}")

        # Загружаем параметры из JSON, если есть
        params_path = os.path.join(company_dir, "image_params.json")
        if os.path.exists(params_path):
            try:
                with open(params_path, 'r', encoding='utf-8') as f:
                    params = json.load(f)

                if 'signature' in params:
                    sig = params['signature']
                    if 'width' in sig:
                        self.signature_image_params["width"] = sig['width']
                    if 'height' in sig:
                        self.signature_image_params["height"] = sig['height']
                    if 'offset_x' in sig:
                        self.signature_image_params["offset_x"] = sig['offset_x']
                    if 'offset_y' in sig:
                        self.signature_image_params["offset_y"] = sig['offset_y']

                if 'stamp' in params:
                    st = params['stamp']
                    if 'width' in st:
                        self.stamp_image_params["width"] = st['width']
                    if 'height' in st:
                        self.stamp_image_params["height"] = st['height']
                    if 'offset_x' in st:
                        self.stamp_image_params["offset_x"] = st['offset_x']
                    if 'offset_y' in st:
                        self.stamp_image_params["offset_y"] = st['offset_y']

                logging.info(f"Загружены параметры изображений для компании {company_name}")
            except Exception as e:
                logging.error(f"Ошибка при загрузке параметров изображений: {e}")

    # ---------------------------------------------------------
    # Конвертация изображения в плавающее
    # ---------------------------------------------------------
    def convert_to_floating_image(self, inline_shape, offset_x, offset_y):
        """Превращает картинку в плавающую 'Перед текстом' с поддержкой X и Y смещения."""
        extent = inline_shape._inline.extent
        graphic = inline_shape._inline.graphic
        docPr = inline_shape._inline.docPr

        anchor_xml = f"""
        <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" 
                   relativeHeight="251658240" behindDoc="0" locked="0" layoutInCell="1" 
                   allowOverlap="1" {nsdecls('wp', 'a', 'pic', 'r')}>
            <wp:simplePos x="0" y="0"/>
            <wp:positionH relativeFrom="character">
                <wp:posOffset>{int(offset_x)}</wp:posOffset>
            </wp:positionH>
            <wp:positionV relativeFrom="line">
                <wp:posOffset>{int(offset_y)}</wp:posOffset>
            </wp:positionV>
            <wp:extent cx="{extent.cx}" cy="{extent.cy}"/>
            <wp:effectExtent l="0" t="0" r="0" b="0"/>
            <wp:wrapNone/>
        </wp:anchor>
        """
        anchor = parse_xml(anchor_xml)
        anchor.append(docPr)
        nv_xml = f'<wp:cNvGraphicFramePr {nsdecls("wp", "a")}><a:graphicFrameLocks noChangeAspect="1"/></wp:cNvGraphicFramePr>'
        anchor.append(parse_xml(nv_xml))
        anchor.append(graphic)
        inline_shape._inline.getparent().replace(inline_shape._inline, anchor)

    # ---------------------------------------------------------
    # Вставка плавающего изображения в run
    # ---------------------------------------------------------
    def insert_floating_image_into_run(self, run, image_params):
        """
        Вставляет плавающее изображение в run с заданными параметрами

        Args:
            run: run, в который будет вставлено изображение
            image_params: словарь с параметрами изображения:
                - path: путь к файлу
                - width: ширина в мм
                - height: высота в мм
                - offset_x: смещение по X в мм
                - offset_y: смещение по Y в мм
        """
        if not image_params.get('path'):
            logging.warning("Путь к изображению не указан")
            return False

        if not os.path.exists(image_params['path']):
            logging.warning(f"Изображение не найдено: {image_params['path']}")
            return False

        try:
            # Сохраняем текст run (плейсхолдер)
            placeholder_text = run.text

            # Очищаем run
            run.text = ""

            # Вставляем изображение с указанными шириной и высотой
            inline_shape = run.add_picture(
                image_params['path'],
                width=Mm(image_params.get('width', 30)),
                height=Mm(image_params.get('height', 30))
            )

            logging.info(
                f"Изображение добавлено в run, размеры: {image_params.get('width', 30)}x{image_params.get('height', 30)} мм")

            # Конвертируем смещения в единицы измерения Word (EMU)
            # 1 мм = 36000 EMU
            offset_x_emu = Mm(image_params.get('offset_x', 0))
            offset_y_emu = Mm(image_params.get('offset_y', -20))

            logging.info(f"Смещение в EMU: X={offset_x_emu}, Y={offset_y_emu}")

            # Превращаем в плавающее изображение
            self.convert_to_floating_image(
                inline_shape,
                offset_x_emu,
                offset_y_emu
            )

            logging.info(f"Изображение успешно вставлено и преобразовано в плавающее")

            return True

        except Exception as e:
            logging.error(f"Ошибка при вставке изображения: {e}")
            import traceback
            logging.error(traceback.format_exc())
            return False

    # ---------------------------------------------------------
    # Вставка подписи и печати (ОБНОВЛЕННЫЙ МЕТОД)
    # ---------------------------------------------------------
    def insert_signature_and_stamp(self, doc):
        """
        Вставляет изображения подписи и печати в документ
        путем замены плейсхолдеров {sign} и {stamp} на изображения
        """
        # Собираем все параграфы (включая таблицы)
        all_paras = list(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_paras.extend(cell.paragraphs)

        # Проверяем наличие изображений
        has_signature = self.signature_image_params.get('path') and os.path.exists(self.signature_image_params['path'])
        has_stamp = self.stamp_image_params.get('path') and os.path.exists(self.stamp_image_params['path'])

        if not has_signature and not has_stamp:
            logging.debug("Нет изображений для вставки")
            return

        signature_count = 0
        stamp_count = 0

        for para in all_paras:
            # Получаем полный текст параграфа
            full_text = "".join(run.text for run in para.runs)

            # Вставляем подпись
            if has_signature and self.SIGNATURE_PLACEHOLDER in full_text:
                for run in list(para.runs):
                    if self.SIGNATURE_PLACEHOLDER in run.text:
                        text_parts = run.text.split(self.SIGNATURE_PLACEHOLDER)
                        run.text = text_parts[0]

                        # Создаем run для изображения
                        img_run = para.add_run(self.SIGNATURE_PLACEHOLDER)

                        # Вставляем изображение
                        if self.insert_floating_image_into_run(img_run, self.signature_image_params):
                            signature_count += 1

                        # Добавляем остаток текста
                        if len(text_parts) > 1 and text_parts[1]:
                            para.add_run(text_parts[1])
                        break

            # Обновляем полный текст для проверки печати
            full_text = "".join(run.text for run in para.runs)

            # Вставляем печать
            if has_stamp and self.STAMP_PLACEHOLDER in full_text:
                for run in list(para.runs):
                    if self.STAMP_PLACEHOLDER in run.text:
                        text_parts = run.text.split(self.STAMP_PLACEHOLDER)
                        run.text = text_parts[0]

                        # Создаем run для изображения
                        img_run = para.add_run(self.STAMP_PLACEHOLDER)

                        # Вставляем изображение
                        if self.insert_floating_image_into_run(img_run, self.stamp_image_params):
                            stamp_count += 1

                        # Добавляем остаток текста
                        if len(text_parts) > 1 and text_parts[1]:
                            para.add_run(text_parts[1])
                        break

        logging.info(f"Вставлено изображений: подпись - {signature_count}, печать - {stamp_count}")

    # ---------------------------------------------------------
    # Совместимость со старым методом _add_signature
    # ---------------------------------------------------------
    def _add_signature(self, docx_path):
        """
        Обертка для совместимости со старым кодом
        Вставляет подпись и печать в документ
        """
        doc = Document(docx_path)
        self.insert_signature_and_stamp(doc)
        doc.save(docx_path)

    # ---------------------------------------------------------
    # JSON
    # ---------------------------------------------------------
    def load_replacements_from_json(self, json_path: str, flatten: bool = True) -> Dict[str, Any]:
        """Загрузка замен из JSON с опциональным уплощением вложенных структур"""
        if not os.path.exists(json_path):
            raise FileNotFoundError(f"JSON файл не найден: {json_path}")

        with open(json_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        if not isinstance(data, dict):
            raise ValueError("JSON должен содержать объект {ключ: значение}")

        if flatten and isinstance(data, dict):
            return self._flatten_dict(data)
        return data

    def _flatten_dict(self, d: Dict, parent_key: str = '') -> Dict:
        """Уплощает вложенный словарь для замен"""
        items = []
        for k, v in d.items():
            new_key = f"{parent_key}.{k}" if parent_key else k
            if isinstance(v, dict):
                items.extend(self._flatten_dict(v, new_key).items())
            else:
                items.append((new_key, v))
        return dict(items)

    # ---------------------------------------------------------
    # Генерация пути
    # ---------------------------------------------------------
    def make_output_path(self, input_path: str, suffix="_updated"):
        folder = os.path.dirname(input_path)
        base, ext = os.path.splitext(os.path.basename(input_path))
        return os.path.join(folder, f"{base}{suffix}{ext}")

    # ---------------------------------------------------------
    # Валидация замен
    # ---------------------------------------------------------
    def validate_replacements(self, doc_path, replacements):
        """Проверяет, какие плейсхолдеры будут заменены"""
        if not os.path.exists(doc_path):
            raise FileNotFoundError(doc_path)

        doc = Document(doc_path)
        found_placeholders = set()
        missing_placeholders = set()

        # Собираем все плейсхолдеры из документа
        doc_text = ""
        for paragraph in doc.paragraphs:
            doc_text += paragraph.text + "\n"

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    doc_text += cell.text + "\n"

        # Проверяем каждый ключ
        for key in replacements.keys():
            if key in doc_text:
                found_placeholders.add(key)
            else:
                # Проверяем вариант со скобками
                if not (key.startswith('{') and key.endswith('}')):
                    if f"{{{key}}}" in doc_text:
                        found_placeholders.add(f"{{{key}}}")
                    else:
                        missing_placeholders.add(key)
                else:
                    missing_placeholders.add(key)

        return {
            'found': list(found_placeholders),
            'missing': list(missing_placeholders),
            'total_found': len(found_placeholders),
            'total_missing': len(missing_placeholders)
        }

    # ---------------------------------------------------------
    # Регулярные выражения для замен
    # ---------------------------------------------------------
    def add_regex_replacement(self, pattern: str, replacement: str, flags=0):
        """Добавляет замену по регулярному выражению"""
        self._regex_replacements.append((re.compile(pattern, flags), replacement))

    def _apply_regex_replacements(self, text: str) -> str:
        """Применяет регулярные замены к тексту"""
        for pattern, replacement in self._regex_replacements:
            text = pattern.sub(replacement, text)
        return text

    # ---------------------------------------------------------
    # Замена в квадратных скобках
    # ---------------------------------------------------------
    def replace_bracket_placeholders(self, text: str, replacements: dict) -> str:
        """
        Заменяет плейсхолдеры в квадратных скобках с поддержкой многострочного текста
        """

        def normalize_text(t):
            """Нормализует текст для сравнения"""
            if not t:
                return ""
            # Удаляем переносы строк и лишние пробелы
            return ' '.join(t.replace('\n', ' ').split())

        def replace_match(match):
            placeholder_raw = match.group(1)
            placeholder_normalized = normalize_text(placeholder_raw)

            # Создаем нормализованные ключи для поиска
            normalized_replacements = {}
            for k, v in replacements.items():
                if k:  # пропускаем пустые ключи
                    normalized_replacements[normalize_text(str(k))] = v

            # 1. Проверяем точное совпадение нормализованного текста
            if placeholder_normalized in normalized_replacements:
                return str(normalized_replacements[placeholder_normalized])

            # 2. Проверяем частичное совпадение
            for norm_key, value in normalized_replacements.items():
                if norm_key and (norm_key in placeholder_normalized or placeholder_normalized in norm_key):
                    return str(value)

            # 3. Проверяем ключевые слова для специфических случаев
            keywords_map = {
                "указать наименование": "organization_name",
                "указать адрес": "organization_address",
                "ПАО": "organization_name",
                "РОСНЕФТЬ": "organization_name"
            }

            for keyword, replace_key in keywords_map.items():
                if keyword in placeholder_normalized:
                    if replace_key in replacements:
                        return str(replacements[replace_key])
                    elif replace_key == "organization_name":
                        return "{zakaz}"
                    elif replace_key == "organization_address":
                        return "{zakaz_adr}"

            # Если замена не найдена, возвращаем текст как был
            return match.group(0)

        # Используем re.DOTALL чтобы точка захватывала переносы строк
        return re.sub(self.BRACKET_PLACEHOLDER_PATTERN, replace_match, text, flags=re.DOTALL)

    def process_paragraph_with_brackets(self, paragraph, replacements):
        # 1. Получаем ПРАВИЛЬНЫЙ полный текст (включая текст внутри гиперссылок)
        full_text = ""
        for child in paragraph._element.xpath('.//w:t'):
            full_text += child.text if child.text else ""

        if '[' not in full_text or ']' not in full_text:
            return

        # 2. Используем логику поиска замен
        new_text = self.replace_bracket_placeholders(full_text, replacements)

        if new_text != full_text:
            # Сохраняем стиль параграфа и формат первого прогона
            first_run_style = None
            if paragraph.runs:
                first_run_style = self._extract_run_formatting(paragraph.runs[0])

            # Очищаем все элементы параграфа (включая гиперссылки)
            p_element = paragraph._p
            for child in list(p_element):
                if child.tag not in (
                        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr',
                        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr'
                ):
                    p_element.remove(child)

            # Добавляем новый чистый текст
            new_run = paragraph.add_run(new_text)
            if first_run_style:
                self._apply_run_formatting(new_run, first_run_style)

    # ---------------------------------------------------------
    # Замена текста с цветным фоном
    # ---------------------------------------------------------
    def replace_colored_text(self, doc, replacements, target_colors=None):
        """
        Заменяет текст определенного цвета
        """
        if target_colors is None:
            target_colors = ['333399']  # Синий цвет в форме Роснефти

        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.font.color and run.font.color.rgb:
                    color = str(run.font.color.rgb).lstrip('#')
                    if color in target_colors:
                        for key, value in replacements.items():
                            if key in run.text:
                                run.text = run.text.replace(key, str(value))

    # ---------------------------------------------------------
    # Основная обработка DOCX
    # ---------------------------------------------------------
    @log_execution
    def replace_text(self, input_path, output_path, replacements, tab_number=1):
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"Входной файл не найден: {input_path}")

        self._log_change('process_start', {
            'input': input_path,
            'output': output_path,
            'replacements_count': len(replacements)
        })

        # Создаем резервную копию
        backup_path = None
        try:
            # Проверяем, доступен ли файл для записи
            if os.path.exists(output_path):
                # Создаем бэкап существующего файла
                backup_path = output_path + '.backup'
                shutil.copy2(output_path, backup_path)

            doc = Document(input_path)

            # Вставляем тег {company_logo}, если его нет после первой строки "Форма ..."
            self._ensure_company_logo_tag(doc)

            # --- СНАЧАЛА ВСЕ ТЕКСТОВЫЕ ЗАМЕНЫ ---
            logging.info("Выполняем текстовые замены...")

            # Выполняем обычные замены (параграфы + таблицы)
            self._process_replacements(doc, replacements)

            # Замена текста с цветным фоном
            self._replace_background_text_in_doc(doc, replacements, partial_match=True)

            # Замена цветного текста (специфика Роснефти)
            self.replace_colored_text(doc, replacements)

            # Удаляем все жёлтые фоны
            self._remove_all_background_colors(doc)

            # --- ПОТОМ ВСТАВЛЯЕМ ПЛАВАЮЩИЕ ИЗОБРАЖЕНИЯ ---
            if tab_number == 2:
                logging.info("Вставляем плавающие изображения подписи и печати...")
                self.insert_signature_and_stamp(doc)

            doc.save(output_path)

            # Удаляем бэкап после успешного сохранения
            if backup_path and os.path.exists(backup_path):
                os.remove(backup_path)

            self._log_change('process_success', {'output': output_path})
            logging.info("DOCX сохранён: %s", output_path)

        except Exception as e:
            # Восстанавливаем из бэкапа при ошибке
            if backup_path and os.path.exists(backup_path):
                shutil.copy2(backup_path, output_path)
                os.remove(backup_path)
            self._log_change('process_error', {'error': str(e)})
            raise RuntimeError(f"Ошибка обработки документа: {e}") from e

    def _replace_background_text_in_doc(self, doc, replacements, partial_match=True):
        """
        Внутренний метод: замена текста с цветным фоном в уже загруженном документе.
        """
        text_with_background = self._find_text_with_background_in_document(doc)

        if not text_with_background:
            return

        for item in text_with_background:
            run_text = item['text']
            run = item['run']

            for old_text, new_text in replacements.items():
                if old_text in run_text:

                    # Сохраняем форматирование
                    font_color = item['font_color']
                    bold = item['bold']
                    italic = item['italic']
                    underline = item['underline']

                    # Замена
                    if partial_match:
                        new_run_text = run_text.replace(old_text, str(new_text))
                    else:
                        if run_text.strip() == old_text:
                            new_run_text = str(new_text)
                        else:
                            continue

                    run.text = new_run_text

                    # Восстанавливаем форматирование
                    if font_color:
                        run.font.color.rgb = font_color
                    run.bold = bold
                    run.italic = italic
                    run.underline = underline

    # ---------------------------------------------------------
    # PDF
    # ---------------------------------------------------------
    @log_execution
    def convert_to_pdf(self, docx_path, pdf_path, add_signature=True):
        if not os.path.exists(docx_path):
            raise FileNotFoundError(docx_path)

        source = docx_path
        if add_signature:
            with self._temp_copy(docx_path) as temp:
                self._add_signature(temp)
                source = temp

        self.converter.convert(source, pdf_path)
        logging.info("PDF создан: %s", pdf_path)

    # ---------------------------------------------------------
    # Пакетная обработка
    # ---------------------------------------------------------
    def batch_process(self, files_config):
        """
        Пакетная обработка нескольких файлов

        files_config: список словарей с ключами:
            - input_path: путь к входному файлу
            - output_path: путь к выходному файлу (опционально)
            - replacements: словарь замен
            - tab_number: номер вкладки (опционально)
            - convert_to_pdf: конвертировать ли в PDF (опционально)
        """
        results = []

        for config in files_config:
            try:
                input_path = config['input_path']
                replacements = config['replacements']

                # Генерация выходного пути если не указан
                if 'output_path' not in config:
                    output_path = self.make_output_path(input_path)
                else:
                    output_path = config['output_path']

                tab_number = config.get('tab_number', 1)

                # Обработка DOCX
                self.replace_text(input_path, output_path, replacements, tab_number)

                # Конвертация в PDF если требуется
                if config.get('convert_to_pdf', False):
                    pdf_path = output_path.replace('.docx', '.pdf')
                    self.convert_to_pdf(output_path, pdf_path,
                                        add_signature=config.get('add_signature', True))
                    results.append({
                        'input': input_path,
                        'docx': output_path,
                        'pdf': pdf_path,
                        'status': 'success'
                    })
                else:
                    results.append({
                        'input': input_path,
                        'docx': output_path,
                        'status': 'success'
                    })

            except Exception as e:
                results.append({
                    'input': config.get('input_path', 'unknown'),
                    'status': 'error',
                    'error': str(e)
                })

        return results

    # ---------------------------------------------------------
    # Контекстный менеджер для работы с документами
    # ---------------------------------------------------------
    @contextmanager
    def open_document(self, path, save_on_exit=True):
        """Контекстный менеджер для безопасной работы с документом"""
        doc = None
        try:
            doc = Document(path)
            yield doc
            if save_on_exit:
                doc.save(path)
        finally:
            pass

    # ---------------------------------------------------------
    # Вставка тега {company_logo} после первой строки "Форма ..."
    # ---------------------------------------------------------
    def _ensure_company_logo_tag(self, doc):
        """
        Ищет ПЕРВУЮ строку 'Форма ...' в верхнем уровне документа
        и вставляет ПОСЛЕ неё новый параграф с тегом {company_logo}
        """

        # Проверяем, есть ли тег вообще
        for p in doc.paragraphs:
            if self.COMPANY_LOGO_PLACEHOLDER in p.text:
                return

        # Ищем ПЕРВУЮ строку, содержащую "Форма"
        target_paragraph = None
        for p in doc.paragraphs:
            if "форма" in p.text.lower():
                target_paragraph = p
                break

        if target_paragraph is None:
            return  # нет формы в верхнем уровне

        # Вставляем новый параграф ПОСЛЕ найденного
        p_elm = target_paragraph._p
        new_p_elm = OxmlElement("w:p")
        p_elm.addnext(new_p_elm)

        new_p = Paragraph(new_p_elm, target_paragraph._parent)
        new_p.text = self.COMPANY_LOGO_PLACEHOLDER
        new_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        logging.info("Вставлен тег {company_logo} после первой строки 'Форма ...'")

    # ---------------------------------------------------------
    # Замены текста и изображений (параграфы + таблицы)
    # ---------------------------------------------------------
    def _process_replacements(self, doc, replacements):
        # Удаляем пустые значения
        replacements = {k: v for k, v in replacements.items() if v not in ("", None)}

        # Нормализуем пути к изображениям
        normalized = {}
        logo_path = None

        for key, val in replacements.items():
            # Проверяем, является ли значение путем к изображению
            if isinstance(val, str) and "." in val:
                # Проверяем существование файла
                if os.path.exists(val):
                    if key in ['company_logo', 'logo', 'Логотип', self.COMPANY_LOGO_PLACEHOLDER]:
                        logo_path = val
                    normalized[key] = val
                elif self.data_dir:
                    possible = os.path.join(self.data_dir, val.replace("/", os.sep))
                    if os.path.exists(possible):
                        if key in ['company_logo', 'logo', 'Логотип', self.COMPANY_LOGO_PLACEHOLDER]:
                            logo_path = possible
                        normalized[key] = possible
                    else:
                        normalized[key] = val
                else:
                    normalized[key] = val
            else:
                normalized[key] = val

        replacements = normalized

        # Если нашли путь к логотипу, добавляем его с правильным ключом
        if logo_path:
            replacements[self.COMPANY_LOGO_PLACEHOLDER] = logo_path

        # Добавляем текущую дату
        self._add_date_replacements(replacements)

        # Расширяем ключи: key → {key}
        extended = {}
        for key, val in replacements.items():
            extended[key] = val
            if not (key.startswith("{") and key.endswith("}")):
                extended[f"{{{key}}}"] = val

        # Добавляем специфические замены
        extended["[указать наименование – ПАО «НК «РОСНЕФТЬ»/ОГ ПАО «НК «РОСНЕФТЬ»]"] = (
            replacements.get("organization_name", replacements.get("{zakaz}", "{zakaz}"))
        )
        extended["[указать адрес ПАО «НК «РОСНЕФТЬ»/ОГ ПАО «НК «РОСНЕФТЬ»]"] = (
            replacements.get("organization_address", replacements.get("{zakaz_adr}", "{zakaz_adr}"))
        )

        # Применяем регулярные выражения
        for key, val in replacements.items():
            if isinstance(val, str):
                replacements[key] = self._apply_regex_replacements(val)

        # Параграфы - НЕ пропускаем плейсхолдеры подписи и печати,
        # так как текстовые замены должны быть выполнены до вставки изображений
        for p in doc.paragraphs:
            self._replace_in_paragraph(p, extended)
            self.process_paragraph_with_brackets(p, extended)
            self._process_hyperlinks(p, extended)

        # Таблицы
        for table in doc.tables:
            self._process_table(table, extended)

    def _add_date_replacements(self, replacements, date=None):
        """Гибкая работа с датами"""
        if date is None:
            date = datetime.date.today()
        elif isinstance(date, str):
            date = datetime.datetime.strptime(date, "%Y-%m-%d").date()

        month_names_ru = [
            "января", "февраля", "марта", "апреля", "мая", "июня",
            "июля", "августа", "сентября", "октября", "ноября", "декабря"
        ]

        month_names_nom = [
            "январь", "февраль", "март", "апрель", "май", "июнь",
            "июль", "август", "сентябрь", "октябрь", "ноябрь", "декабрь"
        ]

        date_replacements = {
            "{day}": f"{date.day:02d}",
            "{day_nopad}": str(date.day),
            "{month}": month_names_ru[date.month - 1],
            "{month_nom}": month_names_nom[date.month - 1],
            "{month_num}": f"{date.month:02d}",
            "{month_num_nopad}": str(date.month),
            "{year}": str(date.year),
            "{year_short}": str(date.year)[-2:],
            "{month_name}": month_names_ru[date.month - 1],
            # Без скобок
            "day": f"{date.day:02d}",
            "month": month_names_ru[date.month - 1],
            "month_num": f"{date.month:02d}",
            "year": str(date.year),
            "month_name": month_names_ru[date.month - 1],
        }

        replacements.update(date_replacements)
        return replacements

    # -------- таблицы с сохранением выравнивания --------
    def _process_table(self, table, replacements):
        """Обработка таблицы с защитой от дублирования и сохранением выравнивания"""
        processed_cells = set()

        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                if (row_idx, col_idx) in processed_cells:
                    continue

                # Сохраняем выравнивание для каждого параграфа в ячейке
                original_alignments = []
                for p in cell.paragraphs:
                    original_alignments.append(p.alignment)

                text = cell.text.strip()
                if not text:
                    continue

                has_brackets = '[' in text and ']' in text
                should_replace = any(k in text for k in replacements.keys()) or has_brackets

                if not should_replace:
                    continue

                # Обработка изображений
                image_replaced = False
                for key, val in replacements.items():
                    if key in text and isinstance(val, str) and os.path.exists(val):
                        self._replace_cell_with_image(cell, key, val)
                        processed_cells.add((row_idx, col_idx))
                        image_replaced = True
                        break

                if image_replaced:
                    continue

                # Текстовая замена
                new_text = text

                # Обычные замены
                for key, val in replacements.items():
                    if key in new_text and not (isinstance(val, str) and os.path.exists(val)):
                        new_text = new_text.replace(key, str(val))

                # Замены в квадратных скобках
                if has_brackets:
                    new_text = self.replace_bracket_placeholders(new_text, replacements)

                if new_text != text:
                    self._clean_cell_content(cell)
                    DocxFormatter.replace_text(cell, new_text, preserve_formatting=True)

                    # Восстанавливаем выравнивание для каждого параграфа
                    for i, p in enumerate(cell.paragraphs):
                        if i < len(original_alignments) and original_alignments[i] is not None:
                            p.alignment = original_alignments[i]

                processed_cells.add((row_idx, col_idx))

    def _clean_cell_content(self, cell):
        """Очищает содержимое ячейки, предотвращая дублирование"""
        while len(cell.paragraphs) > 1:
            p = cell.paragraphs[-1]
            p._p.getparent().remove(p._p)

        if cell.paragraphs:
            p = cell.paragraphs[0]
            while len(p.runs) > 1:
                run = p.runs[-1]
                run._r.getparent().remove(run._r)

            if p.runs:
                p.runs[0].text = ""

    # ---------------------------------------------------------
    # Методы для работы с текстом, выделенным цветом фона
    # ---------------------------------------------------------
    def _iter_runs_with_background(self, doc):
        """Генератор для обхода всех run с фоном в документе"""
        # Параграфы
        for paragraph in doc.paragraphs:
            for run in self._find_runs_with_background_color(paragraph):
                yield run

        # Таблицы
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in self._find_runs_with_background_color(paragraph):
                            yield run

    def _find_runs_with_background_color(self, paragraph):
        runs_with_background = []

        for run in paragraph.runs:
            rPr = run._r.rPr
            if rPr is not None:
                shd = rPr.find(
                    './/w:shd',
                    namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                )
                if shd is not None:
                    fill = shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                    if fill and fill.upper() not in ('FFFFFF', 'AUTO'):
                        runs_with_background.append(run)

        return runs_with_background

    def _find_text_with_background_in_document(self, doc):
        """Находит весь текст с цветным фоном в документе"""
        text_with_background = []

        # Параграфы
        for p_idx, paragraph in enumerate(doc.paragraphs):
            runs = self._find_runs_with_background_color(paragraph)
            for run in runs:
                if run.text.strip():
                    text_with_background.append({
                        'type': 'paragraph',
                        'paragraph_idx': p_idx,
                        'run': run,
                        'text': run.text,
                        'font_color': run.font.color.rgb if run.font.color else None,
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline
                    })

        # Таблицы
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    for p_idx, paragraph in enumerate(cell.paragraphs):
                        runs = self._find_runs_with_background_color(paragraph)
                        for run in runs:
                            if run.text.strip():
                                text_with_background.append({
                                    'type': 'table',
                                    'table_idx': t_idx,
                                    'row_idx': r_idx,
                                    'cell_idx': c_idx,
                                    'paragraph_idx': p_idx,
                                    'run': run,
                                    'text': run.text,
                                    'font_color': run.font.color.rgb if run.font.color else None,
                                    'bold': run.bold,
                                    'italic': run.italic,
                                    'underline': run.underline
                                })

        return text_with_background

    def _remove_all_background_colors(self, doc):
        """Удаляет фон (shading) у всего текста в документе"""
        for run in self._iter_runs_with_background(doc):
            rPr = run._r.rPr
            if rPr is not None:
                shd = rPr.find(
                    './/w:shd',
                    namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                )
                if shd is not None:
                    rPr.remove(shd)

    # ---------------------------------------------------------
    # Методы для работы с изображениями и отступами
    # ---------------------------------------------------------
    def _get_image_dimensions(self, img_path):
        """Кэширование размеров изображений с ограничением размера кэша"""
        if img_path in self._image_cache:
            return self._image_cache[img_path]

        try:
            with Image.open(img_path) as im:
                result = {
                    'ratio': im.height / im.width if im.width else 1.0,
                    'width': im.width,
                    'height': im.height
                }
        except Exception as e:
            logging.warning(f"Не удалось прочитать изображение {img_path}: {e}")
            result = {'ratio': 1.0, 'width': 100, 'height': 100}

        # Управление размером кэша
        if len(self._image_cache) >= self.MAX_CACHE_SIZE:
            self._image_cache.popitem(last=False)
        self._image_cache[img_path] = result

        return result

    def _remove_paragraph_spacing(self, paragraph):
        """Удаляет отступы и интервалы у параграфа с изображением"""
        p_pr = paragraph._p.get_or_add_pPr()

        # Устанавливаем нулевые отступы
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '0')
        spacing.set(qn('w:after'), '0')
        spacing.set(qn('w:line'), '240')  # 1.0 в терминах Word (240 = 1.0)
        spacing.set(qn('w:lineRule'), 'auto')

        # Удаляем существующий spacing если есть
        old_spacing = p_pr.find(qn('w:spacing'))
        if old_spacing is not None:
            p_pr.remove(old_spacing)

        p_pr.append(spacing)

        # Устанавливаем нулевые отступы слева/справа
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '0')
        ind.set(qn('w:right'), '0')
        ind.set(qn('w:firstLine'), '0')
        ind.set(qn('w:hanging'), '0')

        old_ind = p_pr.find(qn('w:ind'))
        if old_ind is not None:
            p_pr.remove(old_ind)

        p_pr.append(ind)

        # Устанавливаем интервал перед/после абзаца в 0
        p_pr.set(qn('w:spaceBefore'), '0')
        p_pr.set(qn('w:spaceAfter'), '0')

    def _add_image_to_paragraph(self, paragraph, img_path, width_inches, height_inches=None):
        """Добавляет изображение в параграф с удалением отступов"""
        try:
            if height_inches is None:
                img_info = self._get_image_dimensions(img_path)
                height_inches = width_inches * img_info['ratio']

            # Добавляем изображение
            run = paragraph.add_run()
            run.add_picture(img_path, width=Inches(width_inches), height=Inches(height_inches))

            # Удаляем отступы у параграфа
            self._remove_paragraph_spacing(paragraph)

            # Также удаляем отступы у всех run в параграфе
            for run in paragraph.runs:
                r_pr = run._r.get_or_add_rPr()
                # Удаляем отступы шрифта если есть
                spacing = r_pr.find(qn('w:spacing'))
                if spacing is not None:
                    r_pr.remove(spacing)

            return True
        except Exception as e:
            logging.warning(f"Не удалось вставить изображение {img_path}: {e}")
            return False

    # ---------------------------------------------------------
    # Параграфы с сохранением выравнивания
    # ---------------------------------------------------------
    def _replace_in_paragraph(self, paragraph, replacements):
        if not paragraph.runs:
            return

        # Сохраняем исходное выравнивание параграфа
        original_alignment = paragraph.alignment

        # Собираем полный текст параграфа
        full_text = "".join(run.text for run in paragraph.runs)

        # Проверяем, есть ли плейсхолдеры для замены
        has_replacements = False
        images_to_insert = []

        # Специальная обработка для company_logo
        if self.COMPANY_LOGO_PLACEHOLDER in full_text:
            logo_path = replacements.get(self.COMPANY_LOGO_PLACEHOLDER)
            if logo_path and isinstance(logo_path, str) and os.path.exists(logo_path):
                has_replacements = True
                images_to_insert.append((self.COMPANY_LOGO_PLACEHOLDER, logo_path))

        # Обычные замены
        for key, val in replacements.items():
            if key in full_text and key != self.COMPANY_LOGO_PLACEHOLDER:
                has_replacements = True
                if isinstance(val, str) and os.path.exists(val):
                    images_to_insert.append((key, val))

        if not has_replacements:
            return

        # Для обычных текстовых замен
        if not images_to_insert:
            new_text = full_text
            for key, val in replacements.items():
                if key in new_text:
                    if isinstance(val, str) and os.path.exists(val):
                        continue
                    new_text = new_text.replace(key, str(val))

            if new_text != full_text:
                self._set_paragraph_text_preserving_format(paragraph, new_text)
                # Восстанавливаем выравнивание
                paragraph.alignment = original_alignment

        # Для замен, включающих изображения
        else:
            # Если в параграфе только один тег
            if len(images_to_insert) == 1 and full_text.strip() == images_to_insert[0][0]:
                key, img_path = images_to_insert[0]

                # Очищаем параграф
                paragraph.clear()

                try:
                    # Определяем размер изображения в зависимости от типа
                    if key == self.COMPANY_LOGO_PLACEHOLDER:
                        # Для логотипа - на всю ширину
                        try:
                            section = paragraph._parent.part.document.sections[0]
                            page_width = section.page_width - section.left_margin - section.right_margin
                            width_inches = page_width / self.EMUS_PER_INCH
                            width_inches = min(width_inches, self.MAX_LOGO_WIDTH)
                        except Exception:
                            width_inches = self.MAX_LOGO_WIDTH

                        with Image.open(img_path) as im:
                            ratio = im.height / im.width if im.width else 1.0

                            margin_inches = self.LOGO_MARGIN
                            final_width = width_inches - (2 * margin_inches)

                            if final_width < 1.0:
                                final_width = width_inches * 0.8

                            # Вставляем изображение с удалением отступов
                            self._add_image_to_paragraph(paragraph, img_path, final_width)

                    else:
                        # Для обычных изображений - стандартный размер
                        self._add_image_to_paragraph(paragraph, img_path, self.DEFAULT_IMAGE_WIDTH)

                    # Восстанавливаем исходное выравнивание
                    paragraph.alignment = original_alignment

                except Exception as e:
                    logging.warning(f"Не удалось вставить изображение {img_path}: {e}")
                    paragraph.add_run(f"[Изображение]")
                    paragraph.alignment = original_alignment
                return

            # Для других одиночных изображений
            elif len(images_to_insert) == 1:
                key, img_path = images_to_insert[0]
                paragraph.clear()

                # Вставляем изображение со стандартным размером и без отступов
                self._add_image_to_paragraph(paragraph, img_path, self.DEFAULT_IMAGE_WIDTH)

                # Восстанавливаем исходное выравнивание
                paragraph.alignment = original_alignment
                return

            # Сложный случай: текст с изображениями внутри
            runs_info = self._extract_runs_info(paragraph)
            paragraph.clear()

            for run_info in runs_info:
                run_text = run_info['text']
                remaining_text = run_text

                while remaining_text:
                    next_image_pos = None
                    next_image_key = None
                    next_image_path = None

                    for key, img_path in images_to_insert:
                        pos = remaining_text.find(key)
                        if pos != -1 and (next_image_pos is None or pos < next_image_pos):
                            next_image_pos = pos
                            next_image_key = key
                            next_image_path = img_path

                    if next_image_pos is not None:
                        if next_image_pos > 0:
                            text_before = remaining_text[:next_image_pos]
                            if text_before:
                                for key, val in replacements.items():
                                    if isinstance(val, str) and os.path.exists(val):
                                        continue
                                    text_before = text_before.replace(key, str(val))

                                new_run = paragraph.add_run(text_before)
                                self._apply_run_formatting(new_run, run_info)

                        # Вставляем изображение с удалением отступов
                        self._add_image_to_paragraph(paragraph, next_image_path, 1.5)

                        remaining_text = remaining_text[next_image_pos + len(next_image_key):]

                    else:
                        for key, val in replacements.items():
                            if isinstance(val, str) and os.path.exists(val):
                                continue
                            remaining_text = remaining_text.replace(key, str(val))

                        if remaining_text:
                            new_run = paragraph.add_run(remaining_text)
                            self._apply_run_formatting(new_run, run_info)
                        break

            # Восстанавливаем исходное выравнивание для всего параграфа
            paragraph.alignment = original_alignment

    def _set_paragraph_text_preserving_format(self, paragraph, new_text):
        if not paragraph.runs:
            paragraph.add_run(new_text)
            return

        first_run = paragraph.runs[0]
        formatting = self._extract_run_formatting(first_run)

        paragraph.clear()
        new_run = paragraph.add_run(new_text)
        self._apply_run_formatting(new_run, formatting)

    def _process_hyperlinks(self, paragraph, replacements):
        """Обрабатывает гиперссылки в параграфе"""
        if not hasattr(paragraph, 'hyperlinks'):
            return

        for hyperlink in paragraph.hyperlinks:
            for run in hyperlink.runs:
                new_text = run.text
                for key, val in replacements.items():
                    if key in new_text:
                        new_text = new_text.replace(key, str(val))

                if new_text != run.text:
                    run.text = new_text

    # ---------------------------------------------------------
    # Обработка изображений в ячейках с сохранением выравнивания
    # ---------------------------------------------------------
    def _replace_cell_with_image(self, cell, placeholder, img_path):
        """Заменяет содержимое ячейки на изображение с сохранением выравнивания"""
        # Сохраняем выравнивание первого параграфа, если есть
        original_alignment = None
        if cell.paragraphs:
            original_alignment = cell.paragraphs[0].alignment

        # Очищаем ячейку
        for p in cell.paragraphs:
            p.clear() if hasattr(p, "clear") else None
        cell.text = ""

        # Получаем или создаем параграф
        if not cell.paragraphs:
            p = cell.add_paragraph()
        else:
            p = cell.paragraphs[0]

        # Удаляем отступы у параграфа
        self._remove_paragraph_spacing(p)

        # Восстанавливаем выравнивание или устанавливаем по центру по умолчанию
        if original_alignment is not None:
            p.alignment = original_alignment
        else:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Получаем ширину колонки
        tc = cell._tc
        tr = tc.getparent()
        tbl = tr.getparent()
        grid = tbl.tblGrid
        col_idx = list(tr).index(tc)

        try:
            col_width_emu = int(grid.gridCol[col_idx].w)
            col_width_inches = col_width_emu / self.EMUS_PER_INCH
            # Оставляем небольшой отступ от краев ячейки
            image_width = col_width_inches * 0.9
        except Exception:
            image_width = 1.5

        # Вставляем изображение
        self._add_image_to_paragraph(p, img_path, image_width)

        # Удаляем отступы у ячейки
        tc_pr = tc.get_or_add_tcPr()
        margins = OxmlElement('w:tcMar')
        for direction in ['top', 'left', 'bottom', 'right']:
            margin = OxmlElement(f'w:{direction}')
            margin.set(qn('w:w'), '0')
            margin.set(qn('w:type'), 'dxa')
            margins.append(margin)

        old_margins = tc_pr.find(qn('w:tcMar'))
        if old_margins is not None:
            tc_pr.remove(old_margins)
        tc_pr.append(margins)

    def _extract_run_formatting(self, run) -> Dict[str, Any]:
        """Извлекает форматирование из run"""
        return {
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline,
            'font': {
                'name': run.font.name,
                'size': run.font.size,
                'color': run.font.color.rgb if run.font.color else None
            }
        }

    def _apply_run_formatting(self, run, formatting: Dict[str, Any]):
        """Применяет форматирование к run"""
        if formatting.get('bold') is not None:
            run.bold = formatting['bold']
        if formatting.get('italic') is not None:
            run.italic = formatting['italic']
        if formatting.get('underline') is not None:
            run.underline = formatting['underline']

        font = formatting.get('font', {})
        if font.get('name'):
            run.font.name = font['name']
        if font.get('size'):
            run.font.size = font['size']
        if font.get('color'):
            run.font.color.rgb = font['color']

    def _extract_runs_info(self, paragraph) -> List[Dict[str, Any]]:
        """Извлекает информацию о форматировании всех runs в параграфе"""
        runs_data = []
        for run in paragraph.runs:
            runs_data.append({
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font': {
                    'name': run.font.name,
                    'size': run.font.size,
                    'color': run.font.color.rgb if run.font.color else None
                }
            })
        return runs_data

    # ---------------------------------------------------------
    # Временная копия
    # ---------------------------------------------------------
    @contextmanager
    def _temp_copy(self, path):
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        tmp.close()
        shutil.copy2(path, tmp.name)
        try:
            yield tmp.name
        finally:
            try:
                os.unlink(tmp.name)
            except OSError:
                logging.warning("Не удалось удалить временный файл: %s", tmp.name)

    # ---------------------------------------------------------
    # Методы для исправления дублирования в таблицах
    # ---------------------------------------------------------
    def fix_duplicated_tables(self, input_path, output_path):
        """
        Исправляет дублирование текста в таблицах существующего документа
        """
        if not os.path.exists(input_path):
            raise FileNotFoundError(input_path)

        doc = Document(input_path)
        fixes_applied = 0

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if self._has_duplicated_content(cell):
                        self._deduplicate_cell(cell)
                        fixes_applied += 1

        if fixes_applied > 0:
            print(f"Исправлено {fixes_applied} ячеек с дублированием")

        doc.save(output_path)
        return output_path

    def _has_duplicated_content(self, cell):
        """Проверяет, есть ли дублирование в ячейке"""
        if len(cell.paragraphs) <= 1:
            return False

        full_text = ""
        for p in cell.paragraphs:
            full_text += p.text

        paragraphs_text = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
        if len(paragraphs_text) < 2:
            return False

        unique_texts = set(paragraphs_text)
        return len(unique_texts) < len(paragraphs_text)

    def _deduplicate_cell(self, cell):
        """Удаляет дубликаты из ячейки"""
        seen_texts = set()
        unique_paragraphs = []

        for p in cell.paragraphs:
            text = p.text.strip()
            if text and text not in seen_texts:
                seen_texts.add(text)
                unique_paragraphs.append(p)

        while cell.paragraphs:
            p = cell.paragraphs[0]
            p._p.getparent().remove(p._p)

        for p in unique_paragraphs:
            new_p = cell.add_paragraph(p.text)
            if p.runs:
                first_run = p.runs[0]
                new_run = new_p.runs[0]
                new_run.bold = first_run.bold
                new_run.italic = first_run.italic
                new_run.underline = first_run.underline

    def _log_change(self, change_type: str, details: Dict[str, Any]):
        """Логирует изменение в истории"""
        self._history.append({
            'timestamp': datetime.datetime.now().isoformat(),
            'type': change_type,
            'details': details
        })

    def get_history(self) -> List[Dict]:
        """Возвращает историю изменений"""
        return self._history.copy()

    def save_history(self, path: str):
        """Сохраняет историю в JSON"""
        with open(path, 'w', encoding='utf-8') as f:
            json.dump(self._history, f, ensure_ascii=False, indent=2)